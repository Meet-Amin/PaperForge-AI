import os
import shutil
import uuid
from io import BytesIO
from pathlib import Path
from typing import Callable, List, Optional, Tuple

import streamlit as st
from dotenv import load_dotenv


ENV_PATH = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=ENV_PATH)


if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError(
        "OPENAI_API_KEY not found.\n"
        f"Expected .env at: {ENV_PATH}\n\n"
        "Put this inside .env (no quotes, no spaces):\n"
        "OPENAI_API_KEY=sk-xxxxxxxxxxxxxxxxxxxxxxxx\n"
    )

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
)

# -----------------------------
# Settings
# -----------------------------
MODEL_NAME = "gpt-4o"
EMBED_MODEL = "text-embedding-3-small"
DEFAULT_K = 6

st.set_page_config(page_title="Research Paper Writer (RAG)", layout="wide")

st.markdown(
    """
<style>
:root {
  --bg-soft: #0b1220;
  --ink: #e7edf7;
  --muted: #9fb0c8;
  --brand: #14b8a6;
  --brand-dark: #0f8f82;
  --card: #111b2e;
  --line: #25324a;
}

.stApp {
  background:
    radial-gradient(circle at 10% 10%, #0f3a3e 0%, transparent 40%),
    radial-gradient(circle at 90% 0%, #1e2d52 0%, transparent 38%),
    var(--bg-soft);
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #0e1628 0%, #0d1422 100%);
}

.hero {
  border: 1px solid var(--line);
  border-radius: 18px;
  padding: 22px 24px;
  background: linear-gradient(135deg, #101a2f 0%, #172744 55%, #103234 100%);
  margin-bottom: 14px;
}

.hero h1 {
  margin: 0 0 8px 0;
  color: var(--ink);
  font-size: 2rem;
  font-family: "Avenir Next", "Gill Sans", "Trebuchet MS", sans-serif;
  letter-spacing: 0.2px;
}

.hero p {
  margin: 0;
  color: var(--muted);
  font-size: 1rem;
}

.section-note {
  color: var(--muted);
  margin-top: -4px;
}

div.stButton > button {
  background: linear-gradient(180deg, var(--brand) 0%, var(--brand-dark) 100%);
  color: white;
  border: none;
  border-radius: 12px;
  padding: 0.6rem 1rem;
  font-weight: 600;
}

div.stDownloadButton > button {
  border-radius: 10px;
}

[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] h1,
[data-testid="stMarkdownContainer"] h2,
[data-testid="stMarkdownContainer"] h3 {
  color: var(--ink);
}
</style>
""",
    unsafe_allow_html=True,
)


# -----------------------------
# Helpers: Save uploads + Load docs
# -----------------------------
def save_uploads_to_folder(uploaded_files, folder: str) -> List[str]:
    os.makedirs(folder, exist_ok=True)
    saved_paths = []
    for f in uploaded_files:
        path = os.path.join(folder, f.name)
        with open(path, "wb") as out:
            out.write(f.getbuffer())
        saved_paths.append(path)
    return saved_paths


def clear_previous_session_data() -> None:
    """Clear persisted artifacts from previous app sessions."""
    db_root = Path("db")
    uploads_root = Path("uploads")

    if db_root.exists():
        for p in db_root.glob("chroma_user_*"):
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)

    if uploads_root.exists():
        for p in uploads_root.iterdir():
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)


def load_docs_from_paths(paths: List[str]) -> List[Document]:
    """
    Loads documents and ensures metadata["source"] contains the filename.
    PyPDFLoader includes page numbers in metadata["page"].
    """
    docs: List[Document] = []

    for p in paths:
        ext = os.path.splitext(p)[1].lower()

        if ext == ".pdf":
            loader = PyPDFLoader(p)
            loaded = loader.load()
            # ensure source is filename
            for d in loaded:
                d.metadata = d.metadata or {}
                d.metadata["source"] = os.path.basename(p)
            docs.extend(loaded)

        elif ext in [".docx", ".doc"]:
            loader = UnstructuredWordDocumentLoader(p)
            loaded = loader.load()
            for d in loaded:
                d.metadata = d.metadata or {}
                d.metadata["source"] = os.path.basename(p)
            docs.extend(loaded)

        elif ext in [".txt", ".md"]:
            loader = TextLoader(p, encoding="utf-8")
            loaded = loader.load()
            for d in loaded:
                d.metadata = d.metadata or {}
                d.metadata["source"] = os.path.basename(p)
            docs.extend(loaded)

        else:
            # Skip unsupported types
            continue

    return docs


def split_docs(
    docs: List[Document], chunk_size: int, chunk_overlap: int
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_documents(docs)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError(
            "DOCX export requires python-docx. Install with: pip install python-docx"
        ) from e

    doc = DocxDocument()
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError as e:
        raise RuntimeError(
            "PDF export requires reportlab. Install with: pip install reportlab"
        ) from e

    out = BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    width, height = A4
    x_margin = 50
    y = height - 50
    line_height = 14

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            y -= line_height
        else:
            if line.startswith("# "):
                c.setFont("Helvetica-Bold", 16)
                text = line[2:].strip()
            elif line.startswith("## "):
                c.setFont("Helvetica-Bold", 13)
                text = line[3:].strip()
            else:
                c.setFont("Helvetica", 11)
                text = line

            max_chars = 100
            wrapped = [
                text[i : i + max_chars] for i in range(0, len(text), max_chars)
            ] or [""]
            for chunk in wrapped:
                if y < 60:
                    c.showPage()
                    y = height - 50
                c.drawString(x_margin, y, chunk)
                y -= line_height

        if y < 60:
            c.showPage()
            y = height - 50

    c.save()
    out.seek(0)
    return out.getvalue()


def render_export_buttons(paper_md: str) -> None:
    st.markdown("### Export Your Paper")
    download_col1, download_col2, download_col3 = st.columns(3)
    with download_col1:
        st.download_button(
            label="Download Markdown",
            data=paper_md.encode("utf-8"),
            file_name="research_paper.md",
            mime="text/markdown",
            use_container_width=True,
        )

    with download_col2:
        try:
            docx_bytes = markdown_to_docx_bytes(paper_md)
            st.download_button(
                label="Download DOCX",
                data=docx_bytes,
                file_name="research_paper.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except RuntimeError as e:
            st.caption(str(e))

    with download_col3:
        try:
            pdf_bytes = markdown_to_pdf_bytes(paper_md)
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name="research_paper.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except RuntimeError as e:
            st.caption(str(e))


# -----------------------------
# Vector DB (Chroma)
# -----------------------------
def build_vector_db(chunks: List[Document], persist_dir: str) -> Chroma:
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)
    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=persist_dir,
        collection_metadata={"hnsw:space": "cosine"},
    )
    return vectordb


def open_vector_db(persist_dir: str) -> Chroma:
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)
    return Chroma(persist_directory=persist_dir, embedding_function=embeddings)


# -----------------------------
# RAG + Citations
# -----------------------------
def format_citation(meta: dict) -> str:
    src = os.path.basename(str(meta.get("source", "uploaded")))
    page = meta.get("page", None)
    if page is not None:
        # PDF pages often 0-indexed internally
        return f"[{src}, p.{int(page) + 1}]"
    return f"[{src}]"


def retrieve_with_citations(db: Chroma, query: str, k: int) -> Tuple[str, List[str]]:
    retriever = db.as_retriever(search_kwargs={"k": k})
    docs = retriever.invoke(query) or []

    if not docs:
        return "", []

    citations: List[str] = []
    seen = set()
    context_parts: List[str] = []

    for d in docs:
        cite = format_citation(d.metadata or {})
        if cite not in seen:
            citations.append(cite)
            seen.add(cite)
        # attach a cite tag the model can reuse
        context_parts.append(f"{d.page_content}\nCITE: {cite}")

    return "\n\n---\n\n".join(context_parts), citations


def llm_generate_section(
    llm: ChatOpenAI,
    section_name: str,
    paper_title: str,
    research_question: str,
    context: str,
) -> str:
    prompt = f"""
You are writing a research paper using ONLY the provided context.

Paper Title: {paper_title}
Research Question: {research_question}

Write the section: {section_name}

Rules:
- Use ONLY the context.
- When you use a fact/claim from the context, include citations inline using the provided CITE tags (e.g., [file.pdf, p.3]).
- Do NOT invent sources, authors, or references.
- If context is insufficient, clearly say what is missing and write the best possible draft anyway.
- Keep it professional and academic friendly.

Context:
{context}
""".strip()

    return llm.invoke(prompt).content.strip()


def generate_full_paper(
    db: Chroma,
    paper_title: str,
    research_question: str,
    k: int,
    progress_callback: Optional[Callable[[float, str], None]] = None,
) -> str:
    llm = ChatOpenAI(model=MODEL_NAME, temperature=0.2)

    sections = [
        "Abstract",
        "Keywords (5-8 keywords)",
        "Introduction",
        "Problem Statement",
        "Related Work / Literature Review",
        "Methodology (use ONLY document info; if missing, propose but label as proposed)",
        "Findings / Analysis (only from documents)",
        "Discussion (interpret findings; keep grounded in docs)",
        "Limitations",
        "Conclusion",
        "Future Work",
    ]

    paper_parts = [f"# {paper_title}\n"]

    total_sections = len(sections)
    for idx, s in enumerate(sections, start=1):
        if progress_callback:
            progress_callback(
                idx / total_sections,
                f"Thinking and writing section {idx}/{total_sections}: {s}",
            )
        query = f"{paper_title}. {research_question}. Write the {s}. Pull relevant details, definitions, methods, and evidence."
        context, _ = retrieve_with_citations(db, query, k=k)
        section_text = llm_generate_section(
            llm, s, paper_title, research_question, context
        )
        paper_parts.append(f"## {s}\n{section_text}\n")

    # References: list unique citations observed in a final retrieval sweep
    _, cites = retrieve_with_citations(
        db,
        f"{paper_title}. {research_question}. What sources were used? Provide citations.",
        k=max(12, k),
    )
    refs = (
        "\n".join([f"- {c}" for c in sorted(set(cites))])
        if cites
        else "- (No citations available from retrieved chunks)"
    )
    paper_parts.append(f"## References\n{refs}\n")

    if progress_callback:
        progress_callback(1.0, "Finalizing paper and references")

    return "\n".join(paper_parts)


def format_paper_for_publication(
    draft_md: str,
    paper_title: str,
    research_question: str,
    author_names: str,
    affiliation: str,
    publication_style: str,
) -> str:
    llm = ChatOpenAI(model=MODEL_NAME, temperature=0.1)
    prompt = f"""
You are an academic editor. Convert the draft paper into a publication-ready manuscript format.

Paper title: {paper_title}
Research question: {research_question}
Author names: {author_names}
Affiliation: {affiliation}
Target style: {publication_style}

Rules:
- Keep claims grounded in the draft only.
- Keep inline citations; do not invent references.
- Improve clarity, flow, and academic tone.
- If author/affiliation is missing, use placeholders.
- Use this structure in Markdown:
  # Title
  Author line
  Affiliation line
  ## Abstract
  ## Keywords
  ## 1. Introduction
  ## 2. Problem Statement
  ## 3. Related Work
  ## 4. Methodology
  ## 5. Findings and Analysis
  ## 6. Discussion
  ## 7. Limitations
  ## 8. Conclusion
  ## 9. Future Work
  ## References

Draft paper:
{draft_md}
""".strip()
    return llm.invoke(prompt).content.strip()


def generate_mandatory_diagram_section(
    paper_md: str,
    paper_title: str,
    research_question: str,
) -> str:
    llm = ChatOpenAI(model=MODEL_NAME, temperature=0.1)
    prompt = f"""
You are an academic assistant.
Create a diagram section for the paper.

Paper title: {paper_title}
Research question: {research_question}

Rules:
- Always return a Markdown section in this format:
  ## 10. Diagram
  <1-2 lines explaining the diagram>
  ```mermaid
  flowchart TD
  ...
  ```
- Use only information implied by the paper draft.
- Keep it simple and readable.

Paper draft:
{paper_md}
""".strip()
    result = llm.invoke(prompt).content.strip()
    # Hard fallback to guarantee a diagram section is always present.
    if "```mermaid" not in result:
        return """
## 10. Diagram
The following flow summarizes the paper development process from evidence to conclusions.
```mermaid
flowchart TD
  A[Uploaded Documents] --> B[Chunking and Embeddings]
  B --> C[Evidence Retrieval]
  C --> D[Section Drafting]
  D --> E[Publication Formatting]
  E --> F[Final Manuscript]
```
""".strip()
    return result


def answer_question_from_docs(
    db: Chroma, question: str, paper_title: str, research_question: str, k: int
) -> Tuple[str, List[str]]:
    context, citations = retrieve_with_citations(
        db,
        f"{paper_title}. {research_question}. User question: {question}",
        k=k,
    )
    if not context:
        return (
            "I could not find relevant evidence in the uploaded documents for that question.",
            [],
        )

    llm = ChatOpenAI(model=MODEL_NAME, temperature=0.1)
    prompt = f"""
Answer the user question using ONLY the context below.

Paper Title: {paper_title}
Research Question: {research_question}
User Question: {question}

Rules:
- Use only the provided context.
- Include inline citations from the given CITE tags.
- If evidence is weak or missing, say that clearly.
- Keep the answer concise and factual.

Context:
{context}
""".strip()
    answer = llm.invoke(prompt).content.strip()
    return answer, citations


# -----------------------------
# Streamlit UI
# -----------------------------
st.markdown(
    """
<div class="hero">
  <h1>PaperForge-AI📚</h1>
  <p>AI Research Paper Generator</p>
</div>
""",
    unsafe_allow_html=True,
)

# On a fresh browser session (including refresh/new tab), clear old persisted artifacts.
if "app_initialized" not in st.session_state:
    clear_previous_session_data()
    st.session_state["app_initialized"] = True

with st.sidebar:
    st.header("Generation Controls")
    st.caption("Tune retrieval and chunking based on document length and complexity.")
    chunk_size = st.slider("Chunk size", 50, 2000, 1000, 50)
    chunk_overlap = st.slider("Chunk overlap", 0, 400, 150, 10)
    top_k = st.slider("Top-K retrieval", 2, 12, DEFAULT_K, 1)
    st.markdown("---")
    st.subheader("Publication Format")
    author_names = st.text_input("Author names", value="Your Name")
    affiliation = st.text_input("Affiliation", value="Your Institution")
    publication_style = st.selectbox(
        "Target style", ["General Academic", "IEEE-like", "APA-like"], index=0
    )
    publication_ready = st.toggle("Make publication-ready format", value=True)
    st.caption("Diagram section is mandatory and will always be included.")

col1, col2 = st.columns([1, 1.15], gap="large")

with col1:
    st.subheader("Step 1: Upload Source Documents")
    st.markdown(
        '<p class="section-note">Supported formats: PDF, DOCX, DOC, TXT, MD.</p>',
        unsafe_allow_html=True,
    )
    uploaded = st.file_uploader(
        "Choose files",
        type=["pdf", "docx", "doc", "txt", "md"],
        accept_multiple_files=True,
    )

    # session-scoped dirs
    session_id = st.session_state.get("session_id") or str(uuid.uuid4())[:8]
    st.session_state["session_id"] = session_id

    persist_dir = st.session_state.get("persist_dir") or f"db/chroma_user_{session_id}"
    st.session_state["persist_dir"] = persist_dir

    if uploaded:
        st.success(f"{len(uploaded)} file(s) ready for ingest.")
    elif st.session_state.get("db_ready") or os.path.exists(persist_dir):
        st.info(f"Using existing vector DB: {persist_dir}")
    else:
        st.warning("No files uploaded yet.")

with col2:
    st.subheader("Step 2: Generate Your Paper")
    st.markdown(
        '<p class="section-note">Define title and objective, then run one click to ingest + write.</p>',
        unsafe_allow_html=True,
    )
    paper_title = st.text_input(
        "Paper title",
        value="A Research Paper Based on Uploaded Documents",
        placeholder="Enter a specific and descriptive title",
    )
    research_question = st.text_area(
        "Research question or goal",
        value="Summarize, analyze, and synthesize the key ideas from the uploaded documents.",
        height=120,
        placeholder="What should the paper answer, compare, or evaluate?",
    )

    if st.button("Generate Paper", use_container_width=True):
        progress_bar = st.progress(0, text="Starting...")

        # Ingest first if files are uploaded in this run
        if uploaded:
            progress_bar.progress(10, text="Saving uploaded files...")
            with st.spinner("Saving uploads..."):
                upload_folder = f"uploads/{session_id}"
                paths = save_uploads_to_folder(uploaded, upload_folder)

            progress_bar.progress(22, text="Loading uploaded documents...")
            with st.spinner("Loading documents..."):
                docs = load_docs_from_paths(paths)

            if not docs:
                st.error("No supported docs loaded (check file types).")
                st.stop()

            progress_bar.progress(35, text="Splitting documents into chunks...")
            with st.spinner("Splitting into chunks..."):
                chunks = split_docs(
                    docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap
                )

            progress_bar.progress(
                50, text="Generating vectors and building Chroma DB..."
            )
            with st.spinner("Creating Chroma vector DB..."):
                _ = build_vector_db(chunks, persist_dir=persist_dir)

            st.success("Documents ingested successfully.")
            st.session_state["db_ready"] = True
        else:
            # Allow re-opening an existing persisted DB when no new upload is provided
            if not st.session_state.get("db_ready"):
                if os.path.exists(st.session_state["persist_dir"]):
                    st.session_state["db_ready"] = True
                else:
                    st.error("Please upload documents first.")
                    st.stop()
            progress_bar.progress(50, text="Using existing vector DB...")

        progress_bar.progress(60, text="Opening vector database...")
        with st.spinner("Opening vector DB..."):
            db = open_vector_db(st.session_state["persist_dir"])

        def on_paper_progress(section_ratio: float, message: str) -> None:
            # Map section progress (0..1) into overall progress range (65..95)
            overall = 65 + int(section_ratio * 30)
            progress_bar.progress(overall, text=message)

        progress_bar.progress(65, text="Writing paper... thinking through sections...")
        with st.spinner("Generating paper (RAG)..."):
            paper_md = generate_full_paper(
                db,
                paper_title,
                research_question,
                k=top_k,
                progress_callback=on_paper_progress,
            )

        if publication_ready:
            progress_bar.progress(96, text="Formatting manuscript for publication...")
            with st.spinner("Applying publication-ready formatting..."):
                paper_md = format_paper_for_publication(
                    draft_md=paper_md,
                    paper_title=paper_title,
                    research_question=research_question,
                    author_names=author_names,
                    affiliation=affiliation,
                    publication_style=publication_style,
                )

        progress_bar.progress(98, text="Generating mandatory diagram section...")
        with st.spinner("Generating diagram section..."):
            diagram_section = generate_mandatory_diagram_section(
                paper_md=paper_md,
                paper_title=paper_title,
                research_question=research_question,
            )
            paper_md = f"{paper_md}\n\n{diagram_section}\n"

        progress_bar.progress(100, text="Paper generation complete")
        st.success("Done!")
        st.session_state["paper_md"] = paper_md
        st.session_state["db_ready"] = True

paper_md_saved = st.session_state.get("paper_md")
if paper_md_saved:
    st.markdown("---")
    render_export_buttons(paper_md_saved)
    st.markdown("---")
    st.markdown(paper_md_saved)

st.markdown("---")
st.subheader("Ask Questions About Your Documents")
st.caption(
    "Chat with your uploaded sources. Answers are grounded in retrieved evidence."
)

if "qa_history" not in st.session_state:
    st.session_state["qa_history"] = []

chat_col1, chat_col2 = st.columns([1, 6])
with chat_col1:
    if st.button("Clear Chat"):
        st.session_state["qa_history"] = []
        st.rerun()

for msg in st.session_state["qa_history"]:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg["role"] == "assistant" and msg.get("citations"):
            st.caption("Citations: " + ", ".join(sorted(set(msg["citations"]))))

question = st.chat_input(
    "Example: What are the key findings and which source supports each one?"
)

if question:
    clean_q = question.strip()
    if clean_q:
        st.session_state["qa_history"].append({"role": "user", "content": clean_q})
        with st.chat_message("user"):
            st.markdown(clean_q)

        if not st.session_state.get("db_ready"):
            if os.path.exists(st.session_state.get("persist_dir", "")):
                st.session_state["db_ready"] = True
            else:
                st.error("Please upload and generate a paper first.")
                st.stop()

        with st.chat_message("assistant"):
            with st.spinner("Searching documents and preparing answer..."):
                db = open_vector_db(st.session_state["persist_dir"])
                answer, cites = answer_question_from_docs(
                    db=db,
                    question=clean_q,
                    paper_title=paper_title,
                    research_question=research_question,
                    k=top_k,
                )
            st.markdown(answer)
            if cites:
                st.caption("Citations: " + ", ".join(sorted(set(cites))))

        st.session_state["qa_history"].append(
            {"role": "assistant", "content": answer, "citations": cites}
        )

st.caption(
    "Tip: For large documents, increase chunk size and Top-K to improve section depth and citation coverage."
)
st.caption(
    
    "© 2026 Meet Amin"
)
